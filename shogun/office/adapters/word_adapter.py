"""Word Adapter — Hybrid python-docx + COM automation for Word documents.

Uses python-docx for all document manipulation (cross-platform, no Office needed).
Uses COM (via com_thread_pool) only for PDF export.

Placeholder convention: ``{{key}}`` — standard mustache style.
"""

from __future__ import annotations

import logging
import re
import time
from pathlib import Path
from typing import Any

log = logging.getLogger("shogun.office.adapters.word")


# ── Document Handle ──────────────────────────────────────────────────


class WordDocumentHandle:
    """Tracks an open Word document (python-docx-based)."""

    def __init__(self, path: Path, document: Any):
        self.path = path
        self.document = document
        self.opened_at = time.time()


# ── Core Adapter Functions ───────────────────────────────────────────


def open_document(file_path: str) -> WordDocumentHandle:
    """Open a Word document with python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for Word operations. Install with: pip install python-docx")

    path = Path(file_path)
    try:
        doc = Document(str(path))
    except Exception as exc:
        error_str = str(exc).lower()
        if "password" in error_str or "encrypted" in error_str:
            from shogun.office.exceptions import PasswordProtectedError
            raise PasswordProtectedError(str(path))
        if "corrupt" in error_str or "not a zip" in error_str or "package not found" in error_str:
            from shogun.office.exceptions import CorruptedFileError
            raise CorruptedFileError(str(path), str(exc))
        raise

    log.debug("Opened document: %s (%d paragraphs)", path, len(doc.paragraphs))
    return WordDocumentHandle(path=path, document=doc)


def close_document(handle: WordDocumentHandle) -> None:
    """Close a Word document (cleanup reference)."""
    handle.document = None
    log.debug("Closed document: %s", handle.path)


def read_text(handle: WordDocumentHandle) -> str:
    """Read the full text content of the document."""
    doc = handle.document
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def read_pages(
    handle: WordDocumentHandle,
    start_page: int = 1,
    end_page: int = 1,
) -> dict[str, Any]:
    """Read a bounded page range using Word's rendered-page markers.

    Word stores ``w:lastRenderedPageBreak`` markers in DOCX files after it
    paginates a document.  Reading those markers avoids sending an entire
    large document to the model when only a few pages were requested.
    """
    if start_page < 1:
        raise ValueError("start_page must be at least 1.")
    if end_page < start_page:
        raise ValueError("end_page must be greater than or equal to start_page.")

    from docx.oxml.ns import qn

    pages: list[list[str]] = [[]]
    body = handle.document.element.body

    def append(value: str) -> None:
        pages[-1].append(value)

    def walk(element: Any) -> None:
        tag = element.tag
        if tag == qn("w:lastRenderedPageBreak"):
            pages.append([])
            return
        if tag == qn("w:br") and element.get(qn("w:type")) == "page":
            pages.append([])
            return
        if tag == qn("w:t"):
            append(element.text or "")
            return
        if tag == qn("w:tab"):
            append("\t")
            return
        if tag == qn("w:br"):
            append("\n")
            return

        for child in element:
            walk(child)

        if tag == qn("w:p"):
            append("\n")
        elif tag == qn("w:tc"):
            append("\t")
        elif tag == qn("w:tr"):
            append("\n")

    walk(body)
    page_count = len(pages)
    if start_page > page_count:
        raise ValueError(
            f"start_page {start_page} exceeds the document's {page_count} rendered pages."
        )

    actual_end = min(end_page, page_count)
    text = "".join(
        part
        for page in pages[start_page - 1:actual_end]
        for part in page
    ).strip()
    return {
        "text": text,
        "length": len(text),
        "start_page": start_page,
        "end_page": actual_end,
        "page_count": page_count,
    }


def read_headings(handle: WordDocumentHandle) -> list[dict[str, Any]]:
    """Read all headings from the document.

    Returns a list of dicts with 'level', 'text', and 'index'.
    """
    doc = handle.document
    headings = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 0
            headings.append({
                "level": level,
                "text": para.text,
                "index": i,
            })
    return headings


def replace_placeholders(
    handle: WordDocumentHandle,
    mapping: dict[str, str],
) -> dict[str, int]:
    """Replace {{placeholder}} patterns in the document.

    Searches paragraphs, tables, headers, and footers.

    Args:
        handle: The document handle.
        mapping: Dict of placeholder → replacement value.
                 Keys should include the braces: ``{{company_name}}``.

    Returns:
        Dict of placeholder → replacement count.
    """
    doc = handle.document
    counts: dict[str, int] = {k: 0 for k in mapping}

    def replace_in_paragraph(paragraph):
        for key, value in mapping.items():
            if key in paragraph.text:
                # Replace in runs to preserve formatting
                full_text = paragraph.text
                if key in full_text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            counts[key] += 1
                    # If runs didn't catch it (split across runs), rebuild
                    if key in paragraph.text:
                        # Fallback: merge and re-split
                        inline = paragraph.runs
                        combined = "".join([r.text for r in inline])
                        if key in combined:
                            new_text = combined.replace(key, value)
                            if inline:
                                inline[0].text = new_text
                                for r in inline[1:]:
                                    r.text = ""
                                counts[key] += 1

    # Paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    # Headers and Footers
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer and hasattr(header_footer, 'paragraphs'):
                for para in header_footer.paragraphs:
                    replace_in_paragraph(para)

    log.debug("Replaced placeholders: %s", {k: v for k, v in counts.items() if v > 0})
    return counts


def insert_paragraph(
    handle: WordDocumentHandle,
    text: str,
    style: str = "Normal",
) -> None:
    """Append a new paragraph to the document."""
    doc = handle.document
    doc.add_paragraph(text, style=style)
    log.debug("Inserted paragraph (style=%s): %s...", style, text[:50])


def create_document_from_text(
    output_path: str,
    text: str,
    append: bool = False,
) -> WordDocumentHandle:
    """Create or append to a Word document and return its open handle."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for Word operations. Install with: pip install python-docx")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(out)) if append and out.exists() else Document()
    lines = text.splitlines()
    if lines:
        for line in lines:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph(text)
    doc.save(str(out))
    log.info("%s document text: %s (%d chars)", "Appended" if append else "Created", out, len(text))
    return WordDocumentHandle(path=out, document=doc)


def insert_table(
    handle: WordDocumentHandle,
    headers: list[str],
    rows: list[list[Any]],
    style: str = "Table Grid",
) -> None:
    """Insert a table into the document.

    Args:
        handle: The document handle.
        headers: Column header strings.
        rows: 2D list of cell values.
        style: Word table style name.
    """
    doc = handle.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))

    try:
        table.style = style
    except Exception:
        pass  # Style may not exist in template

    # Headers
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            if col_idx < len(headers):
                table.rows[row_idx + 1].cells[col_idx].text = str(value) if value is not None else ""

    log.debug("Inserted table: %d columns × %d rows", len(headers), len(rows))


def apply_style(
    handle: WordDocumentHandle,
    paragraph_index: int,
    style_name: str,
) -> None:
    """Apply a named style to a paragraph by index."""
    doc = handle.document
    if paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"Paragraph index {paragraph_index} out of range (max {len(doc.paragraphs) - 1}).")
    doc.paragraphs[paragraph_index].style = style_name


def save_as(handle: WordDocumentHandle, output_path: str) -> str:
    """Save the document to a new file path."""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    handle.document.save(str(out))
    log.info("Saved document to: %s", out)
    return str(out)


def get_document_metadata(handle: WordDocumentHandle) -> dict[str, Any]:
    """Get metadata about the document."""
    doc = handle.document
    props = doc.core_properties
    return {
        "file": str(handle.path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "title": props.title or "",
        "author": props.author or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "last_modified_by": props.last_modified_by or "",
    }


# ── COM-Only Functions (PDF Export) ──────────────────────────────────


def _com_export_pdf(file_path: str, output_path: str, visible: bool = False) -> str:
    """Export a Word document to PDF via COM.

    Must be called on the STA thread pool.
    """
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = visible
    word.DisplayAlerts = 0  # wdAlertsNone
    out = str(Path(output_path).resolve())
    try:
        doc = word.Documents.Open(str(Path(file_path).resolve()))
        try:
            # wdFormatPDF = 17
            doc.SaveAs2(out, FileFormat=17)
        finally:
            doc.Close(SaveChanges=0)
    finally:
        word.Quit()

    log.info("Exported document to PDF: %s", out)
    return out


async def export_pdf(file_path: str, output_path: str, visible: bool = False) -> str:
    """Export to PDF — async wrapper around COM call."""
    from shogun.office.com_thread_pool import run_com, office_lock
    async with office_lock("word"):
        return await run_com(_com_export_pdf, file_path, output_path, visible)
