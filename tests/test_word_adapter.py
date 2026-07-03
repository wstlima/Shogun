from docx import Document
from docx.oxml import OxmlElement

from shogun.office.adapters.word_adapter import (
    create_document_from_text,
    open_document,
    read_pages,
    read_text,
)
from shogun.services.native_skills import NATIVE_TOOLS, generate_tool_prompt


def _add_rendered_page_break(paragraph) -> None:
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:lastRenderedPageBreak"))


def test_read_pages_returns_only_requested_rendered_pages(tmp_path):
    path = tmp_path / "rendered-pages.docx"
    document = Document()
    first = document.add_paragraph("Page one")
    _add_rendered_page_break(first)
    second = document.add_paragraph("Page two")
    _add_rendered_page_break(second)
    document.add_paragraph("Page three")
    document.save(path)

    result = read_pages(open_document(str(path)), 1, 2)

    assert result["page_count"] == 3
    assert result["start_page"] == 1
    assert result["end_page"] == 2
    assert "Page one" in result["text"]
    assert "Page two" in result["text"]
    assert "Page three" not in result["text"]


def test_read_pages_rejects_invalid_range(tmp_path):
    path = tmp_path / "one-page.docx"
    Document().save(path)

    handle = open_document(str(path))

    try:
        read_pages(handle, 2, 2)
    except ValueError as exc:
        assert "exceeds" in str(exc)
    else:
        raise AssertionError("Expected an out-of-range page request to fail.")


def test_read_pages_tool_is_exposed_with_json_prompt_format():
    tools = {
        tool["function"]["name"]: tool
        for tool in NATIVE_TOOLS
    }

    assert "office_word_read_pages" in tools
    prompt = generate_tool_prompt([tools["office_word_read_pages"]])
    assert "valid JSON" in prompt
    assert 'tool_name({"param1": "value1"' in prompt


def test_create_document_from_text_writes_complete_content(tmp_path):
    path = tmp_path / "translated.docx"
    translated = "Første afsnit\nAndet afsnit med danske tegn: æøå"

    handle = create_document_from_text(str(path), translated)

    assert path.exists()
    assert read_text(handle) == translated


def test_create_document_from_text_can_append(tmp_path):
    path = tmp_path / "translated.docx"
    create_document_from_text(str(path), "Side et")

    handle = create_document_from_text(str(path), "Side to", append=True)

    assert read_text(handle) == "Side et\nSide to"


def test_create_from_text_tool_is_exposed():
    names = {
        tool["function"]["name"]
        for tool in NATIVE_TOOLS
    }

    assert "office_word_create_from_text" in names
    assert "office_word_read_page" in names
